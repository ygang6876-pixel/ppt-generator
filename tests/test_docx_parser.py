from pathlib import Path

from docx import Document
from PIL import Image

from src.docx_parser import parse_docx


def test_parse_docx_keeps_tables_and_images_near_captions(tmp_path: Path):
    image_path = tmp_path / "site.png"
    Image.new("RGB", (120, 80), (220, 120, 40)).save(image_path)

    docx_path = tmp_path / "scheme.docx"
    document = Document()
    document.add_paragraph("施工专项方案")
    document.add_heading("1 工程概况", level=1)
    document.add_paragraph("本页用于测试施工方案正文提取。")
    document.add_paragraph("图1-1 施工平面布置图", style="Heading 6")
    document.add_picture(str(image_path))
    document.add_paragraph("表1-1 主要工程量表", style="Heading 6")
    table = document.add_table(rows=2, cols=2)
    table.rows[0].cells[0].text = "项目"
    table.rows[0].cells[1].text = "数量"
    table.rows[1].cells[0].text = "锚杆"
    table.rows[1].cells[1].text = "120"
    document.save(docx_path)

    deck = parse_docx(docx_path)

    assert deck.metadata["theme"] == "construction"
    assert any(slide.images for slide in deck.slides)
    assert any(slide.tables for slide in deck.slides)
    assert any(slide.title == "图1-1 施工平面布置图" for slide in deck.slides)
    assert any(slide.title == "表1-1 主要工程量表" for slide in deck.slides)
    assert not any(slide.title == "目 录" for slide in deck.slides)
