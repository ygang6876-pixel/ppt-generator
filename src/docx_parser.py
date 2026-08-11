from __future__ import annotations

from pathlib import Path
import re

from docx import Document
from docx.document import Document as DocxDocument
from docx.oxml.ns import qn
from docx.table import Table as DocxTable
from docx.text.paragraph import Paragraph

from .markdown_parser import Deck, ImageAsset, Slide, Table


MAX_TEXT_ITEMS_PER_SLIDE = 6
CAPTION_STYLE_NAMES = {"caption", "题注"}
BRIEF_TRIGGER_SLIDE_COUNT = 45
BRIEF_MAX_SLIDES = 36

REPORT_SECTIONS = [
    ("工程概况", r"工程概况|工程特性|周边环境"),
    ("施工部署", r"施工平面|施工道路|施工供电|施工供水|资源投入|施工进度"),
    ("工艺技术", r"施工工艺|工艺流程|截水沟|防护网|土方|石方|爆破|支护|锚杆|喷混|套拱|框格"),
    ("风险控制", r"风险|危险|监测|安全|质量|环保|雨季|冬季"),
    ("验收应急", r"验收|应急|救援|强制性标准"),
]

PRIORITY_RULES = [
    (r"工程概况|工程特性", 2),
    (r"主要工程量|工程量表", 2),
    (r"周边环境|施工平面|施工道路|布置图|示意图", 5),
    (r"施工要求|目标", 2),
    (r"风险辨识|安全风险|危险源|预防措施", 3),
    (r"施工进度|资源投入|设备投入|劳动力", 4),
    (r"主要施工原则|工艺流程", 2),
    (r"截水沟|防护网|土方开挖|石方明挖|爆破|起爆|支护|锚杆|喷混|套拱|框格", 12),
    (r"安全监测|安全保障|质量|文明施工|环境保护|水土保持|雨季|冬季", 7),
    (r"验收|应急|救援|强制性标准", 4),
]

REPORT_GROUPS = [
    (
        6,
        [
            r"工程概况|工程特性",
            r"主要工程量|工程量表",
            r"周边环境|施工平面|布置图|示意图",
        ],
    ),
    (
        4,
        [
            r"施工进度|进度安排",
            r"资源投入|设备投入|劳动力",
            r"施工道路|施工供电|施工供水|施工排水",
        ],
    ),
    (
        11,
        [
            r"主要施工原则|工艺流程",
            r"截水沟",
            r"被动防护网|防护网",
            r"土方开挖",
            r"石方明挖|爆破|起爆",
            r"边坡支护|锚杆|挂网喷混|喷混",
            r"套拱|框格|洞门",
        ],
    ),
    (
        6,
        [
            r"风险辨识|安全风险|危险源|预防措施",
            r"安全监测|监测",
            r"安全保障|安全管理|安全生产",
            r"质量|文明施工|环境保护|水土保持|雨季|冬季",
        ],
    ),
    (3, [r"应急|救援|处置"]),
    (1, [r"强制性标准"]),
    (4, [r"验收"]),
]


def parse_docx(path: Path, mode: str = "brief") -> Deck:
    document = Document(path)
    title = _document_title(document, path)
    subtitle = _document_subtitle(document)
    slides: list[Slide] = []
    current: Slide | None = None
    pending_caption: str | None = None
    media_dir = path.parent / f"{path.stem}_assets"
    image_index = 1

    for block in _iter_blocks(document):
        if isinstance(block, Paragraph):
            text = _clean_text(block.text)
            style_name = block.style.name if block.style is not None else ""

            if _is_toc(style_name):
                continue

            images = _extract_paragraph_images(block, media_dir, image_index)
            image_index += len(images)

            if images and current is None and pending_caption is None:
                continue

            if text and _is_caption(style_name, text):
                pending_caption = text
                if images:
                    slides.append(
                        Slide(
                            title=text,
                            body=[],
                            images=images,
                            layout="image-full",
                        )
                    )
                    pending_caption = None
                continue

            if images:
                image_title = pending_caption or text or (current.title if current else "图片")
                image_body = [] if pending_caption else ([text] if text else [])
                slides.append(
                    Slide(
                        title=image_title,
                        body=image_body[:2],
                        images=images,
                        layout="image-full" if len(images) == 1 else "image-right",
                    )
                )
                pending_caption = None
                continue

            if not text:
                continue

            if _is_section_heading(style_name):
                current = Slide(title=text)
                slides.append(current)
                pending_caption = None
                continue

            if _is_minor_heading(style_name):
                if current is None:
                    current = Slide(title=text)
                    slides.append(current)
                else:
                    _append_text(current, text)
                pending_caption = None
                continue

            if current is None:
                current = Slide(title="内容")
                slides.append(current)

            if _is_list(block):
                _append_bullet(current, text)
            else:
                _append_text(current, text)

        elif isinstance(block, DocxTable):
            table = _convert_table(block)
            if table is None:
                continue

            table_title = pending_caption or "表格"
            slides.append(
                Slide(
                    title=table_title,
                    tables=[table],
                    layout="full-table",
                )
            )
            pending_caption = None

    slides = _remove_empty_slides(slides)
    if mode == "brief" and len(slides) > BRIEF_TRIGGER_SLIDE_COUNT:
        slides = _curate_report_slides(slides)
    if not slides:
        raise ValueError("该文档未读取到可生成 PPT 的正文内容。")

    return Deck(
        title=title,
        subtitle=subtitle,
        slides=slides,
        metadata={"theme": "construction", "footer": "PPT Generator"},
    )


def _iter_blocks(document: DocxDocument):
    body = document.element.body
    for child in body.iterchildren():
        if child.tag == qn("w:p"):
            yield Paragraph(child, document)
        elif child.tag == qn("w:tbl"):
            yield DocxTable(child, document)


def _document_title(document, path: Path) -> str:
    cover_lines = _cover_lines(document)
    for index, line in enumerate(cover_lines):
        if "方案" in line:
            previous = cover_lines[index - 1] if index > 0 else ""
            if previous and not re.search(r"批准|审核|校核|编制|合同|公司|项目部", previous):
                return f"{previous}{line}"
            return line
    if document.core_properties.title and "方案" in document.core_properties.title:
        return document.core_properties.title
    for paragraph in document.paragraphs:
        text = _clean_text(paragraph.text)
        if text and not _is_toc(paragraph.style.name if paragraph.style else ""):
            return text
    return path.stem


def _document_subtitle(document) -> str:
    lines = _cover_lines(document)
    return " / ".join(lines[1:3]) if len(lines) > 1 else "Generated from WPS/Word document"


def _cover_lines(document) -> list[str]:
    lines: list[str] = []
    for paragraph in document.paragraphs:
        text = _clean_text(paragraph.text)
        style_name = paragraph.style.name if paragraph.style is not None else ""
        if not text or _is_toc(style_name):
            continue
        if _is_section_heading(style_name):
            break
        lines.append(text)
        if len(lines) >= 12:
            break
    return lines


def _clean_text(text: str) -> str:
    return re.sub(r"\s+", " ", text or "").strip()


def _is_toc(style_name: str) -> bool:
    return style_name.lower().startswith("toc")


def _heading_level(style_name: str) -> int | None:
    normalized = style_name.lower()
    match = re.match(r"heading\s+(\d+)", normalized)
    if match:
        return int(match.group(1))
    match = re.match(r"标题\s*(\d+)", style_name)
    if match:
        return int(match.group(1))
    return None


def _is_section_heading(style_name: str) -> bool:
    level = _heading_level(style_name)
    return level is not None and level <= 3


def _is_minor_heading(style_name: str) -> bool:
    level = _heading_level(style_name)
    return level is not None and 3 < level < 6


def _is_caption(style_name: str, text: str) -> bool:
    normalized = style_name.lower()
    if normalized in CAPTION_STYLE_NAMES or style_name in CAPTION_STYLE_NAMES:
        return True
    return bool(re.match(r"^(图|表)\s*\d", text)) or _heading_level(style_name) == 6


def _is_list(paragraph) -> bool:
    style_name = paragraph.style.name if paragraph.style is not None else ""
    if "list" in style_name.lower() or "列表" in style_name:
        return True
    text = _clean_text(paragraph.text)
    return bool(re.match(r"^([（(]?\d+[）).、]|[①②③④⑤⑥⑦⑧⑨⑩]|[-•])", text))


def _append_text(slide: Slide, text: str) -> None:
    if not text:
        return
    if len(slide.body) + len(slide.bullets) < MAX_TEXT_ITEMS_PER_SLIDE:
        slide.body.append(_shorten(text))


def _append_bullet(slide: Slide, text: str) -> None:
    if not text:
        return
    if len(slide.body) + len(slide.bullets) < MAX_TEXT_ITEMS_PER_SLIDE:
        slide.bullets.append(_shorten(text))


def _shorten(text: str, limit: int = 120) -> str:
    text = _clean_text(text)
    if len(text) <= limit:
        return text
    return f"{text[:limit].rstrip()}..."


def _extract_paragraph_images(paragraph: Paragraph, media_dir: Path, start_index: int) -> list[ImageAsset]:
    images: list[ImageAsset] = []
    for offset, blip in enumerate(paragraph._element.iter(qn("a:blip"))):
        rel_id = blip.get(qn("r:embed"))
        if not rel_id or rel_id not in paragraph.part.related_parts:
            continue
        image_part = paragraph.part.related_parts[rel_id]
        ext = Path(str(image_part.partname)).suffix or _extension_from_content_type(image_part.content_type)
        media_dir.mkdir(parents=True, exist_ok=True)
        filename = f"docx_image_{start_index + offset:03d}{ext.lower()}"
        image_path = media_dir / filename
        if not image_path.exists():
            image_path.write_bytes(image_part.blob)
        images.append(ImageAsset(alt=paragraph.text.strip() or image_path.stem, path=str(image_path.relative_to(media_dir.parent))))
    return images


def _extension_from_content_type(content_type: str) -> str:
    mapping = {
        "image/png": ".png",
        "image/jpeg": ".jpg",
        "image/jpg": ".jpg",
        "image/gif": ".gif",
        "image/bmp": ".bmp",
        "image/webp": ".webp",
    }
    return mapping.get(content_type, ".png")


def _convert_table(docx_table: DocxTable) -> Table | None:
    rows = [[_clean_text(cell.text) for cell in row.cells] for row in docx_table.rows]
    rows = [row for row in rows if any(cell for cell in row)]
    if not rows:
        return None
    if len(rows) == 1:
        headers = [f"列{index + 1}" for index in range(len(rows[0]))]
        data_rows = rows
    else:
        headers = rows[0]
        data_rows = rows[1:]
    col_count = len(headers)
    return Table(headers=headers, rows=[_normalize_row(row, col_count) for row in data_rows])


def _normalize_row(row: list[str], size: int) -> list[str]:
    if len(row) < size:
        return row + [""] * (size - len(row))
    return row[:size]


def _remove_empty_slides(slides: list[Slide]) -> list[Slide]:
    return [
        slide
        for slide in slides
        if slide.body or slide.bullets or slide.images or slide.tables or slide.mermaid or slide.code_blocks
    ]


def _curate_report_slides(slides: list[Slide]) -> list[Slide]:
    selected: list[Slide] = []
    selected_keys: set[int] = set()
    intro = _build_intro_slide(slides)
    agenda = _build_agenda_slide(slides)
    if intro:
        selected.append(intro)
    if agenda:
        selected.append(agenda)

    for group_limit, patterns in REPORT_GROUPS:
        group_count = 0
        for pattern in patterns:
            for index, slide in enumerate(slides):
                if index in selected_keys or group_count >= group_limit:
                    continue
                if not _matches(slide, pattern):
                    continue
                selected.append(_style_report_slide(slide))
                selected_keys.add(index)
                group_count += 1
                if len(selected) >= BRIEF_MAX_SLIDES:
                    return selected[:BRIEF_MAX_SLIDES]

    for index, slide in enumerate(slides):
        if index in selected_keys:
            continue
        if not (slide.images or slide.tables):
            continue
        selected.append(_style_report_slide(slide))
        if len(selected) >= BRIEF_MAX_SLIDES:
            break

    return selected[:BRIEF_MAX_SLIDES]


def _build_intro_slide(slides: list[Slide]) -> Slide | None:
    facts: list[str] = []
    for slide in slides[:20]:
        for item in [*slide.body, *slide.bullets]:
            if re.search(r"\d", item) or re.search(r"危险性|高边坡|支洞|洞口|坡比|开挖|支护", item):
                facts.append(_shorten(item, 95))
            if len(facts) >= 5:
                break
        if len(facts) >= 5:
            break
    if not facts:
        return None
    return Slide(title="汇报重点", bullets=facts[:5], layout="highlight")


def _build_agenda_slide(slides: list[Slide]) -> Slide | None:
    items = [name for name, pattern in REPORT_SECTIONS if any(_matches(slide, pattern) for slide in slides)]
    if not items:
        return None
    return Slide(title="汇报目录", bullets=items[:5], layout="agenda")


def _matches(slide: Slide, pattern: str) -> bool:
    haystack = " ".join([slide.title, *slide.body[:2], *slide.bullets[:2]])
    return bool(re.search(pattern, haystack))


def _style_report_slide(slide: Slide) -> Slide:
    title = _clean_report_title(slide.title)
    body = [_shorten(item, 95) for item in slide.body[:4]]
    bullets = [_shorten(item, 95) for item in slide.bullets[:5]]
    layout = slide.layout
    if slide.tables:
        layout = "full-table"
    elif slide.images:
        layout = "blueprint" if _looks_like_drawing(title) else ("image-full" if len(slide.images) == 1 else "image-right")
    elif re.search(r"流程|进度|计划|步骤|程序", title):
        layout = "process"
    elif re.search(r"风险|危险|坍塌|触电|打击|伤害|爆炸", title):
        layout = "risk"
    elif re.search(r"措施|目标|职责|保障|应急|验收|标准|管理|控制", title):
        layout = "checklist"
    elif re.search(r"工程概况|工程特性|项目概况|施工部署", title):
        layout = "overview"
    elif len(body) + len(bullets) <= 4:
        layout = "summary"
    return Slide(
        title=title,
        body=body,
        bullets=bullets,
        images=slide.images[:2],
        tables=slide.tables[:1],
        mermaid=slide.mermaid,
        code_blocks=slide.code_blocks,
        layout=layout,
    )


def _clean_report_title(title: str) -> str:
    title = _clean_text(title)
    title = re.sub(r"^\d+(\.\d+)*\s*", "", title)
    return title or "内容"


def _looks_like_drawing(title: str) -> bool:
    return bool(re.search(r"图|布置|示意|断面|设计|流程|网络|参数", title))
